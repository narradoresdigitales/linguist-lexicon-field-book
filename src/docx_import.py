from typing import List, Dict, Any, Tuple, Optional
from datetime import datetime
from docx import Document

def _safe_header(h: str) -> str:
    return (h or "").strip().lower()

def _str(v) -> str:
    return "" if v is None else str(v).strip()

def load_docx(path_or_bytes) -> Document:
    """Accepts a file-like object (BytesIO) or a path."""
    return Document(path_or_bytes)

def extract_tables_as_dicts(doc: Document) -> List[List[Dict[str, Any]]]:
    """
    Return list of tables; each table is a list of dict rows keyed by normalized headers.
    If a table lacks headers, it will be returned as empty (skip it).
    """
    tables_rows: List[List[Dict[str, Any]]] = []
    for tbl in doc.tables:
        if len(tbl.rows) < 2:
            continue
        headers = [_safe_header(cell.text) for cell in tbl.rows[0].cells]
        if not any(headers):
            # no header row detected
            continue
        rows: List[Dict[str, Any]] = []
        for r in tbl.rows[1:]:
            cells = [cell.text for cell in r.cells]
            row = {headers[i] if i < len(headers) else f"col{i}": _str(cells[i]) for i in range(len(cells))}
            rows.append(row)
        if rows:
            tables_rows.append(rows)
    return tables_rows

def extract_plain_text(doc: Document) -> str:
    """Concatenate paragraphs and list items into a single text blob."""
    parts: List[str] = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)
    return "\n".join(parts)

def candidate_words_from_text(text: str) -> List[str]:
    """
    Generate a de-duplicated, case-normalized list of candidate words from free text.
    - Keeps alphabetic and hyphenated terms (e.g., 'state-of-the-art').
    - Filters out very short tokens (len < 2) and numeric tokens.
    """
    import re
    tokens = re.findall(r"[A-Za-z][A-Za-z\-']{1,}", text)
    # normalize: lowercase and strip trailing punctuation-like chars
    normalized = []
    for t in tokens:
        t = t.strip("—-–.,:;!?()[]{}\"' \t").lower()
        if len(t) >= 2 and not t.isnumeric():
            normalized.append(t)
    # de-duplicate while preserving order
    seen = set()
    uniq = []
    for w in normalized:
        if w not in seen:
            seen.add(w)
            uniq.append(w)
    return uniq

def map_row_to_entry(row: Dict[str, Any], default_tags=None, default_source: str = "") -> Dict[str, Any]:
    """
    Map a row from a table into your app's entry schema.
    Supports header variants like 'term' -> 'word'.
    """
    default_tags = default_tags or []
    header_aliases = {
        "word": ["word", "term", "vocabulary", "entry"],
        "definition": ["definition", "meaning", "gloss"],
        "notes": ["notes", "context", "example", "examples"],
        "tags": ["tags", "label", "labels"],
        "source": ["source", "class", "course"],
        "timestamp": ["timestamp", "time", "t"],
    }
    def pick(keys: List[str]) -> str:
        for k in keys:
            if k in row and row[k]:
                return _str(row[k])
        return ""

    word = pick(header_aliases["word"])
    definition = pick(header_aliases["definition"])
    notes = pick(header_aliases["notes"])
    tags_str = pick(header_aliases["tags"])
    source = pick(header_aliases["source"]) or default_source
    timestamp = pick(header_aliases["timestamp"])

    # coerce tags
    tags = default_tags.copy()
    if tags_str:
        # allow CSV-like or bracketed lists
        if tags_str.startswith("["):
            tags.extend([t.strip(" '\"") for t in tags_str.strip("[]").split(",") if t.strip()])
        else:
            tags.extend([t.strip() for t in tags_str.split(",") if t.strip()])
    # normalize & unique tags
    norm = []
    seen = set()
    for t in tags:
        tt = t.strip()
        if tt and tt.lower() not in seen:
            seen.add(tt.lower())
            norm.append(tt)

    return {
        "word": word,
        "definition": definition,
        "notes": notes,
        "tags": norm,
        "source": source,
        "timestamp": timestamp,
        "date_added": datetime.utcnow().isoformat(timespec="seconds") + "Z",
    }